import os
import csv
import base64
import traceback

import PyPDF2
# I wouldn't suggest using PyFPDF, as it is outdated and no longer being maintained.
# Instead, you could use fpdf2 (see the documentation as well)
import docx2txt
import html2text
from docx import Document
from docx.opc import exceptions
import pytesseract
from pdf2image import convert_from_path
from PIL import Image

from src.logger.Logger import Logger


log = Logger().log


class AttachmentParser:
    def __init__(self):
        self.tmp_fldr = f'{os.getcwd()}/temp/'
        os.makedirs(self.tmp_fldr, exist_ok=True)
        self.docx_ext = '.docx'

    def error_processor(self, ex: Exception) -> None:
        """Обработчик сообщений об ошибке"""
        log.error("An exception of type {0} occurred. Arguments:\n{1!r}\n".format(type(ex).__name__, ex.args))
        if type(ex) == exceptions.PackageNotFoundError:
            log.info('CANT PROCESS ATTACHMENT FILE')
        else:
            [log.info(f'{item}') for item in traceback.format_exception(type(ex), ex, ex.__traceback__)]
        return None

    @staticmethod
    def attach_extention_check(attach_name: str) -> bool:
        ext = attach_name.split('.')[-1]
        return ext in ['docx', 'doc', 'pdf', 'xls', 'xlsx', 'rtf']

    @staticmethod
    def attach_name_check(attach_name):
        return 'тз' not in attach_name.lower().split(' ') and 'заяв' not in attach_name.lower()

    def parse_docx(self, attachment_path) -> str:
        """Parser via python-docx"""
        try:
            document = Document(''.join(attachment_path.split('.')[:-1]) + self.docx_ext)
        except KeyError:
            document = None

        full_text = ''
        if document is not None and len(document.tables) > 0:
            for table_count, _ in enumerate(document.tables):
                table = document.tables[table_count]
                for i, row in enumerate(table.rows):
                    text = [cell.text + '\n' for cell in row.cells if len(cell.text) > 2]
                    full_text += '\n '.join(text)

            paragraphs = []
            for num, para in enumerate(document.paragraphs):
                paragraphs.append(para.text)
            full_text += ' '.join(paragraphs)

        if len(full_text) == 0:
            extracted_list = []
            for x in docx2txt.process(''.join(attachment_path.split('.')[:-1]) + self.docx_ext).split():
                if x not in extracted_list:
                    extracted_list.append(x)
            return ' '.join(extracted_list)

        return full_text

    def parse_pdf(self, attachment_path) -> str:
        """Parser via PyPDF2"""
        filename = ''.join(attachment_path.split('.')[:-1]) + '.pdf'
        document = open(filename, 'rb')
        pdf_reader = PyPDF2.PdfReader(document)
        page = pdf_reader.pages[0]
        try:
            full_text = page.extract_text()
        except UnicodeDecodeError as ex:
            full_text = ''
            self.error_processor(ex)

        # OCR parsing
        if len(full_text) == 0:
            try:
                pdf_pages = convert_from_path(filename, 500)
            except Image.DecompressionBombError:
                pdf_pages = []

            image_file_list = []
            for page_enumeration, page in enumerate(pdf_pages, start=1):
                filename = f"{self.tmp_fldr}/page_{page_enumeration:04}.jpg"
                page.save(filename, "JPEG")
                image_file_list.append(filename)

            for image_file in image_file_list:
                full_text = pytesseract.image_to_string(Image.open(image_file), lang="rus")

        return full_text

    @staticmethod
    def parse_csv(attachment_path) -> str:
        """Parser via csvreader"""
        full_text = ''
        with open(''.join(attachment_path.split('.')[:-1]) + '.csv', newline='') as csvfile:
            reader = csv.reader(csvfile, delimiter=',')
            for row in reader:
                full_text += ' ' + ' '.join(row)
        return full_text

    def save_attachments_files(self, parsed_eml) -> list:
        os.makedirs(self.tmp_fldr, exist_ok=True)
        attach_names = []
        if 'attachment' in parsed_eml.keys() and len(parsed_eml['attachment']) > 0:
            for attach_count in range(len(parsed_eml['attachment'])):
                attach_name = parsed_eml['attachment'][attach_count]['filename']
                attach_name = attach_name.replace('.', '')[:40] + "." + attach_name.split(".")[-1]
                if self.attach_extention_check(attach_name) and self.attach_name_check(attach_name):
                    f = open(f'{self.tmp_fldr}/{attach_name}', 'wb+')
                    f.write(base64.b64decode(parsed_eml['attachment'][attach_count]['raw']))
                    f.close()
                    attach_names.append(attach_name)
        log.info(f'Attachment list - {attach_names}')
        return attach_names

    def convert_attachment_file(self, attach_name: str):
        os.makedirs(self.tmp_fldr, exist_ok=True)
        cmd = ''
        attachment_path = f'{self.tmp_fldr}{attach_name}'
        if attach_name.split('.')[-1] == 'doc':
            cmd = f'lowriter --convert-to docx "{attachment_path}" --outdir "{self.tmp_fldr}"'
        elif attach_name.split('.')[-1] == 'xlsx' or attach_name.split('.')[-1] == 'xls':
            cmd = f'unoconv -f csv "{attachment_path}"'
        elif attach_name.split('.')[-1] == 'rtf':
            cmd = f'unoconv -f docx "{attachment_path}"'
        os.system(cmd)

        return attachment_path

    @staticmethod
    def org_structure(inn, bik, r_account, corr_account, email, first, middle, last, tel, website, address):
        return {
            'inn': [inn if len(inn) == 10 or len(inn) == 12 else inn + '*'][0],
            'bik': [bik if len(bik) == 9 and bik[:2] == '04' else bik + '*'][0],
            'r_account': [r_account if len(r_account) == 20 else r_account + '*'][0],
            'c_account': [corr_account if len(corr_account) == 20 else corr_account + '*'][0],
            'email': email,
            "person_fname": first,
            "person_pname": middle,
            "person_sname": last,
            "phone": tel,
            "website": website,
            "address": address,
            "bank": "",
            "kpp": "",
            "ogrn": "",
            "okato": ""
        }

    @staticmethod
    def parse_message_text(parsed_eml: dict):
        """Parse text from html-formatted email"""
        return html2text.html2text(parsed_eml['body'][-1]['content'])

    def get_message_text_and_header(self, parsed_eml: dict) -> (str, str):
        # TO-DO - split into two separate methods
        message_text = ''
        if len(parsed_eml['body']) > 0:
            message_text += self.parse_message_text(parsed_eml)

        texts_to_replace = ['http://www.rusbelt.ru', 'rusbelt']
        for text_to_replace in texts_to_replace:
            message_text = message_text.replace(text_to_replace, " ")

        header_from = parsed_eml['header']['header']['from'][0]

        return message_text, header_from
